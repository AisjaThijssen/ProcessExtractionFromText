import os
import re
from typing import List, Tuple, Union, Optional, Dict, Any

import pandas as pd
import spacy
from docx import Document
from docx.shared import RGBColor
from spacy.tokens.token import Token
from spacy.tokens.doc import Doc
from itertools import combinations
import graphviz

# load spaCy Dutch model
nlp = spacy.load("nl_core_news_sm")


def load_data(path: str) -> pd.DataFrame:
    """Load the data from a txt file splitting the text into Articles.

    Args:
        path (str): The path to where the data is located

    Returns:
        pd.DataFrame: A DataFrame containing the name of the Article and its
            content.
    """
    with open(path, encoding="utf-8") as file:
        document_text = file.read()

    # Split the document at newlines and remove empty newlines.
    document_text = document_text.split("\n")
    document_text = [sentence for sentence in document_text if sentence != ""]

    documents = [[]]
    i = 0

    # Loop through the sentences and when you find an indicator of a new
    # Article start a new entry in documents so we can store all sentences that
    # belong to one Article together.
    for sentence in document_text:
        if bool(re.fullmatch(r"(Artikel \S+)( \([^)]+\))?", sentence)):
            documents.append(
                [re.fullmatch(r"(Artikel \S+)( \([^)]+\))?", sentence).group(1)]
            )
            i += 1
        else:
            documents[i].append(sentence)

    # Remove the first document as it does not contain an Article but any
    # information before the first article.
    documents = documents[1:]

    # Store the content of the Articles.
    content = [
        " ".join(document[1:]).replace("\t", " ")
        for document in documents
        if "[Vervallen per" not in " ".join(document[1:]).replace("\t", " ")
    ]

    # Store the names of the Articles
    titles = [
        document[0]
        for document in documents
        if "[Vervallen per" not in " ".join(document[1:]).replace("\t", " ")
    ]

    # Put everything into a DataFrame
    data = pd.DataFrame({"article": titles, "content": content})

    return data


def extract_sub_sentences(parse_tree: Doc) -> List[List[Token]]:
    """Extract all sub sentences from a main sentence using certain dependency
    tags and the subtrees belonging to the words with those tags.

    Args:
        parse_tree (Doc): A SpaCy parsing tree of a full sentence including
            dependencies that can be used to identify sub-sentences.

    Returns:
        List[List[Token]]: A list with all identified sub-sentences.
    """
    sub_sentences = []

    # Reconstruct the full sentence so we can check that the sub-sentence will
    # not be the same a s the full sentence.
    sentence = "".join(word.text_with_ws for word in parse_tree)

    for word in parse_tree:
        if word.dep_ in ["acl", "acl:relcl", "advcl", "csubj", "ccomp", "xcomp"]:
            sub_sentence = "".join(w.text_with_ws for w in word.subtree)

            # Identify if the sub-sentence is al least 3 words long and if the
            # sub-sentence is really a sub-sentence (shorter than the full
            # sentence).
            if len(sub_sentence.split(" ")) >= 3 and len(sub_sentence) < len(sentence):
                sub_sentences.append([word for word in word.subtree])

    return sub_sentences


def identify_active_sentence(dependencies: List[str]) -> bool:
    """Identify if a sentence is written in active form or in passive form.

    Args:
        dependencies (List[str]): A list of all dependency tags of the words in
            a sentence.

    Returns:
        bool: True if the sentence is in active form and False if the sentence
            is in passive form.
    """
    for dependency in dependencies:
        if "pass" in dependency:
            return False

    return True


def is_copula(word: Token) -> bool:
    """Identify if a word is part of a copula.

    Args:
        word (Token): The word we are looking at.

    Returns:
        bool: True if the words is part of a copula, False otherwise.
    """
    # We say a word is part of a copula if the head is a form of to be and the
    # word itself is its subject.
    return word.head.lemma_ == "zijn" and word.dep_ == "nsubj"


def create_actor(
    actor: Tuple[Token, Optional[str]], tree: Doc, sentence_id: int
) -> Dict[str, Any]:
    """[Algorithm 7] This function creates an actor object from a word that is
    indicated to be an actor.

    Args:
        actor (Tuple[Token, Optional[str]]): The main actor token found by the
            algorithm.
        tree (Doc): A parsing tree of the sentence.
        sentence_id (int): The id of the sentence we are working on.

    Returns:
        Dict[str, Any]: The original token, the modifiers relevant to that
            token, the text representation of the actor, a boolean indicating
            if the actor is the subject of the sentence (always set to True),
            an actor that this actor is referencing to (always set to None),
            and the id of the sentence this actor is in.
    """
    relevant_words = {actor[0]}
    new_words = True

    while new_words:
        # Find all noun specifiers based on the current relevant words and the
        # dependency tags.
        modifiers = {
            word
            for word in tree
            if (
                word.dep_
                in [
                    "nmod:poss",
                    "det",
                    "amod",
                    "iobj",
                    "compound:prt",
                    "xcomp",
                    "nmod",
                    "acl:relcl",
                    "case",
                    "fixed",
                    "appos",
                ]
            )
            # Only include words that are related to a word we have so far but
            # are not in our current list yet.
            and (word.head in relevant_words) and (word not in relevant_words)
        }

        # If no new words are found we stop the loop.
        if not modifiers:
            new_words = False
            continue

        # Otherwise we add the words to the relevant words and repeat.
        else:
            relevant_words.update(modifiers)

    actor_object = {
        "object": actor,
        "relevant_words": relevant_words,
        "text": " ".join(word.text.lower() for word in tree if word in relevant_words),
        "subject": True,
        "reference": None,
        "sentence_id": sentence_id,
    }
    return actor_object


def create_object(
    object: Tuple[Token, Optional[str]], tree: Doc, sentence_id: int
) -> Dict[str, Any]:
    """[Algorithm 8] This function creates an object object from a word that is
    indicated to be an object.

    Args:
        object (Tuple[Token, Optional[str]]): The main object token found by
            the algorithm.
        tree (Doc): A parsing tree of the sentence.
        sentence_id (int): The id of the sentence we are working on.

    Returns:
        Dict[str, Any]: The original token, the modifiers relevant to that
            token, the text representation of the object, a boolean indicating
            if the object is the subject of the sentence (always set to False),
            an actor that this actor is referencing to (always set to None),
            and the id of the sentence this object is in.
    """
    # Create an actor out of the object.
    object_object = create_actor(object, tree, sentence_id)
    # Set subject to False to distinguish between actors and objects.
    object_object["subject"] = False
    return object_object


def create_action(
    action: Tuple[Token, Optional[str]], tree: Doc, sentence_id: int
) -> Dict[str, Any]:
    """[Algorithm 10] This function creates an action object from a word that
    is indicated to be an action.

    Args:
        action (Tuple[Token, Optional[str]]): The main action token found by
            the algorithm.
        tree (Doc): A parsing tree of the sentence.
        sentence_id (int): The id of the sentence we are working on.

    Returns:
        Dict[str, Any]: The original token, the modifiers relevant to that
            token, the text representation of the action, the id of the
            sentence this action is in, a list of markers found fot the action,
            indicators for each marker if they are from a complementizer.
    """
    relevant_words = {action[0]}
    new_words = True

    while new_words:
        # Find all verb specifiers based on the current relevant words and the
        # dependency tags.
        modifiers = {
            word
            for word in tree
            if (
                word.dep_
                in [
                    "aux",
                    "aux:pass",
                    "advcl",
                    "advmod",
                    "cop",
                    "nmod",
                    "acl:relcl",
                    "case",
                    "fixed",
                    "xcomp",
                    "mark",
                    "obl",
                    "det",
                    "appos",
                    "amod",
                ]
            )
            # Only include words that are related to a word we have so far but
            # are not in our current list yet.
            and (word.head in relevant_words) and (word not in relevant_words)
        }

        # If no new words are found we stop the loop.
        if not modifiers:
            new_words = False
            continue

        # Otherwise we add the words to the relevant words and repeat.
        else:
            relevant_words.update(modifiers)

    action_object = {
        "object": action,
        "relevant_words": relevant_words,
        "text": " ".join(word.text.lower() for word in tree if word in relevant_words),
        "sentence_id": sentence_id,
        "markers": [],
        "marker_from_complementizer": [],
    }
    return action_object


def find_dependency(
    dependency: str, tree: Doc, all: bool = False
) -> Union[List[Token], Token, None]:
    """Find the word(s) with a certain dependency tag in a sentence if
    available.

    Args:
        dependency (str): The dependency label we are looking for.
        tree (Doc): A parsing tree of the sentence.
        all (bool, optional): True if we want to find all occurrences of the
            dependency label, False if we only want to find the first one.
            Defaults to False.

    Returns:
        Any[List[Token], Token, None]: A list of word tokens if we want to find
            all occurrences of a dependency tag, one word token if we only want
            to find the first occurrence of a dependency label and None if no
            token with the given tag was found.
    """
    if all:
        dependencies = []
        for word in tree:
            if word.dep_ == dependency:
                dependencies.append(word)
        return dependencies
    else:
        # If all is False we only return the first occurrence of the relevant
        # dependency.
        for word in tree:
            if word.dep_ == dependency:
                return word
    return None


def check_conjunctions(
    active: bool, tree: Doc, element: Token, element_type: str
) -> List[Tuple[Token, str]]:
    """[Algorithm 5] Check if there are any conjunctions to the element given
    as input.

    Args:
        active (bool): True if the sentence is in active form and False if the
            sentence is in passive form.
        tree (Doc): A parsing tree of the sentence.
        element (Token): The token of the word of interest.
        element_type (str): The label of the word of interest (action, actor or
            object).

    Returns:
        List[Tuple[Token, str]]: A list of all conjunctions found with their
            labels.
    """
    conjunctions = []
    # go through all words that are marked as conjunctions in the parsing tree
    dependencies = find_dependency("conj", tree, all=True)
    for dependency in dependencies:
        # Identify if the open clausal complement of the action is the governor
        # of the conjunction.
        xcomp_match = False
        if element_type == "action":
            xcomp = find_dependency("xcomp", tree)
            if xcomp:
                if xcomp == dependency.head:
                    xcomp_match = True

        # If the element has a relation to the conjunction and the element is
        # not part of a copula or when the above condition holds we need to
        # create a new object.
        if (
            (element == dependency.head) and (not is_copula(dependency.head))
        ) or xcomp_match:
            conjunction_node = dependency
            if xcomp_match:
                new_element = element
                new_element_type = "action"

                # The line below is a part of the pseudocode but has no effect
                # on the code here hence it is commented out.

                # new_element_xcomp = conjunction_node
            else:
                if element_type == "action":
                    new_element = conjunction_node
                    new_element_type = "action"
                else:
                    new_element = conjunction_node
                    new_element_type = "object"
            conjunctions.append((new_element, new_element_type))

            # To prevent infinite loops we make sure that the new element is
            # not the same as the considered element.
            if new_element != element:
                conjunctions += check_conjunctions(
                    active, tree, new_element, new_element_type
                )

    return conjunctions


def determine_actors(active: bool, tree: Doc) -> List[Tuple[Token, str]]:
    """[Algorithm 3] Determine the actors present in a sentence.

    Args:
        active (bool): True if the sentence is in active form and False if the
            sentence is in passive form.
        tree (Doc): A parsing tree of the sentence.

    Returns:
        List[Tuple[Token, str]]: A list of all actors found with their labels.
    """
    actors = []
    # In an active sentence we look for the subject and in a passive sentence
    # we look for the agent.
    if active:
        actor = find_dependency("nsubj", tree)
    else:
        actor = find_dependency("obl:agent", tree)

    if actor:
        actors.append((actor, "actor"))
        actors += check_conjunctions(active, tree, actor, "actor")

    return actors


def determine_actions(active: bool, tree: Doc) -> List[Tuple[Token, str]]:
    """[Algorithm 4] Determine the actions present in a sentence.

    Args:
        active (bool): True if the sentence is in active form and False if the
            sentence is in passive form.
        tree (Doc): A parsing tree of the sentence.

    Returns:
        List[Tuple[Token, str]]: A list of all actions found with their labels.
    """
    actions = []
    main_predicate = None
    if active:
        subject = find_dependency("nsubj", tree)
        if subject:
            main_predicate = subject.head

            # In case the main verb is a copula verb we want to make sure that
            # the main_predicate is set to the referent and not the form of to
            # be as this gives more information.
            cop = find_dependency("cop", tree)
            if cop:
                if main_predicate == cop:
                    main_predicate = cop.head

        else:
            object = find_dependency("obj", tree)
            if object:
                main_predicate = object.head
    else:
        subject = find_dependency("nsubj:pass", tree)

        # We noticed in outputs where no nsubj:pass was found that looking for
        # nsubj still have us the correct action.
        if not subject:
            subject = find_dependency("nsubj", tree)

        if subject:
            main_predicate = subject.head

    action = main_predicate
    if action:
        actions.append((action, "action"))
        actions += check_conjunctions(active, tree, action, "action")

    return actions


def determine_objects(
    active: bool, tree: Doc, action: Token
) -> List[Tuple[Token, str]]:
    """[Algorithm 6] Determine the objects present in a sentence, belonging to
    a certain action.

    Args:
        active (bool): True if the sentence is in active form and False if the
            sentence is in passive form.
        tree (Doc): A parsing tree of the sentence.
        action (Token): The actions for which we want the object.

    Returns:
        List[Tuple[Token, str]]: A list of all objects found with their labels.
    """
    objects = []
    object_node = None

    if active:
        object = find_dependency("obj", tree)
        if object:
            object_node = object
        else:
            cop = find_dependency("cop", tree)
            if cop:
                if cop.head.dep_ == "nsubj":
                    object_node = cop.head
    else:
        object_node = find_dependency("nsubj:pass", tree)
        # Similarly to when we were dealing with the action in the passive
        # sentence we use nsubj if not nsubj:pass can be found.
        if not object_node:
            object_node = find_dependency("nsubj", tree)

    if object_node:
        objects.append((object_node, "object"))
        objects += check_conjunctions(active, tree, object_node, "object")

    return objects


def extract_elements(tree: Doc, sentence_id: int) -> List[Dict[str, Any]]:
    """[Algorithm 2] Extract all actor, action and object combinations from a
    sentence.

    Args:
        tree (Doc): A parsing tree of the input sentence.
        sentence_id (int): The id of the sentence we are working on.

    Returns:
        List[Dict[str, Any]]: A list of all extracted combinations of an
            action, actor and object.
    """
    dependencies = [word.dep_ for word in tree if word.is_alpha]
    # Identify if the sentence is active.
    active = identify_active_sentence(dependencies)

    # Find all actors and create actor object from them.
    actors = determine_actors(active, tree)
    actors = [create_actor(actor, tree, sentence_id) for actor in actors]

    # Find all actions.
    raw_actions = determine_actions(active, tree)

    actions_with_object = []

    # If no actions were found we append an empty tuple
    if len(raw_actions) == 0:
        actions_with_object.append((None, None))

    # Loop over the actions and determine the objects for each of them.
    for action in raw_actions:
        objects = determine_objects(active, tree, action)
        if len(objects) > 0:
            # For each object action combination we create a new entry in the
            # actions_with_objects list making the actions and objects into
            # their respective objects.
            for object in objects:
                actions_with_object.append(
                    (
                        create_action(action, tree, sentence_id),
                        create_object(object, tree, sentence_id),
                    )
                )
        else:
            # If no object was found save the action with a none object.
            actions_with_object.append((create_action(action, tree, sentence_id), None))

    final_actions = []
    # Combine each action object pair with each actor found in the sentence.
    for action in actions_with_object:
        if action[0]:
            if len(actors) > 0:
                for actor in actors:
                    final_actions.append(
                        {
                            "action": action[0].copy() if action[0] else None,
                            "actor": actor.copy(),
                            "object": action[1].copy() if action[1] else None,
                            "link": None,
                            "link_type": None,
                        }
                    )
            else:
                final_actions.append(
                    {
                        "action": action[0].copy() if action[0] else None,
                        "actor": None,
                        "object": action[1].copy() if action[1] else None,
                        "link": None,
                        "link_type": None,
                    }
                )

    return final_actions


def sentence_decomposition(
    tree: Doc, sentence_id: int, completed_sentences: List[List[Token]] = []
) -> List[Dict[str, Any]]:
    """[Algorithm 1] Split the sentence in sub-sentences and extract elements
    for each sub-sentence.

    Args:
        tree (Doc): A parsing tree of the input sentence.
        sentence_id (int): The id of the sentence we are working on.
        completed_sentences (List[List[Token]]): A sentences already processed.
            Defaults to an empty list.

    Returns:
        List[Dict[str, Any]]: A list of all extracted combinations of an
            action, actor and object for a sentence.
    """
    extracted_elements = []

    # Find the sub sentences
    sub_sentences = extract_sub_sentences(tree)

    # If there are not sub-sentences find the elements in the sentence.
    if len(sub_sentences) == 0:
        # Avoid infinite loops by only considering new sentences.
        if tree not in completed_sentences:
            completed_sentences.append(tree)
            extracted_elements += extract_elements(tree, sentence_id)

    # If there is a sub-sentence split the sub-sentence from the main sentence
    # and extract elements from both.
    elif len(sub_sentences) == 1:
        extracted_elements += sentence_decomposition(sub_sentences[0], sentence_id)

        # Identify the main sentence.
        main_sentence = [word for word in tree if word not in sub_sentences[0]]
        dependencies = [word.dep_ for word in main_sentence if word.is_alpha]
        if any(
            dep in dependencies for dep in ["nsubj", "nsubj:pass", "obl:agent", "obj"]
        ):
            extracted_elements += sentence_decomposition(main_sentence, sentence_id)

    # If there are more than one sub-sentence extract elements from each
    # sub-sentence and the remaining main sentence.
    else:
        remaining_sentence = tree

        for sub_sentence in sub_sentences:
            extracted_elements += sentence_decomposition(sub_sentence, sentence_id)

            # Identify the words that remain.
            remaining_sentence = [
                word for word in remaining_sentence if word not in sub_sentence
            ]

        dependencies = [word.dep_ for word in remaining_sentence if word.is_alpha]

        if any(
            dep in dependencies for dep in ["nsubj", "nsubj:pass", "obl:agent", "obj"]
        ):
            extracted_elements += sentence_decomposition(
                remaining_sentence, sentence_id, completed_sentences
            )

    return extracted_elements


def choose_best_fit(
    reference_results: List[Tuple[Dict[str, Any], float]]
) -> Dict[str, Any]:
    """Choose the reference that fits best according to the scores. If two
    options are tied, choose the one the closest to the word that is bing
    resolved.

    Args:
        reference_results (List[Tuple[Dict[str, Any], float]]): A list of all
            referencing candidates and their score

    Returns:
        Dict[str, Any]: The best candidate
    """
    # Extract all scores so we can find the best scoring.
    scores = [result[1] for result in reference_results]

    # If no candidates exist return None.
    if len(scores) == 0:
        return None

    max_score = max(scores)
    objects_with_highest_score = [
        result[0] for result in reference_results if result[1] == max_score
    ]

    # If we have multiple results with the highest score choose the one with
    # the highest sentence id, this will be the closest to the word we are
    # resolving.
    if len(objects_with_highest_score) > 1:
        sentence_ids = [object["sentence_id"] for object in objects_with_highest_score]
        max_sentence_id = max(sentence_ids)
        objects_with_highest_sentence_id = [
            object
            for object in objects_with_highest_score
            if object["sentence_id"] == max_sentence_id
        ]
        # In case there are multiple words with the same sentence id we return
        # the last one as it will be closest to the word we are resolving.
        return objects_with_highest_sentence_id[-1]
    else:
        return objects_with_highest_score[0]


def find_reference(
    sentence_id: int,
    object_to_resolve: Dict[str, Any],
    sentences: List[Dict[str, Any]],
    distance_penalty: float = -15,
    role_match_score: float = 20,
    subject_role_score: float = 10,
    object_role_score: float = 10,
):
    """[Algorithm 12] Find the reference for the words that require one.

    Args:
        sentence_id (int): The id of the sentence that the word we are trying
            to resolve is in.
        object_to_resolve (Dict[str, Any]): the object we are trying to
            resolve.
        sentences (List[Dict[str, Any]]): All sentences and their extracted
            events.
        distance_penalty (float, optional): The penalty given to being far way
            from the word we are resolving. Defaults to -15.
        role_match_score (float, optional): The score we give to a word having
            the right role (object or subject). Defaults to 20.
        subject_role_score (float, optional): The score we give to a word if
            the word to be resolved is not part of a copula and this word is
            the actor in its event. Defaults to 10.
        object_role_score (float, optional): The score we give to a word if the
            word to be resolved is part of a copula and this word is the object
            in its event. Defaults to 10.

    Returns:
        Dict[str, Any]: the chosen reference object.
    """
    # we cannot find a reference for a word in the first sentence hence we
    # return None.
    if sentence_id == 0:
        return None

    # Create a list of all potential references out of all actors and objects
    # that occur before the element to tbe resolved.
    potential_references = []
    for sentence in sentences[: sentence_id + 1]:
        for event in sentence["events"]:
            for object in ["actor", "object"]:
                if event[object]:
                    # If it is in the same sentence it must have an index
                    # before the element to be resolved.
                    if event[object]["sentence_id"] == sentence_id:
                        if (
                            event[object]["object"][0].i
                            < object_to_resolve["object"][0].i
                        ):
                            potential_references.append(event[object])
                    else:
                        potential_references.append(event[object])

    results = []

    for object in potential_references:
        # If an object already has a reference look to that reference instead.
        while object.get("reference"):
            if object.get("reference"):
                object = object.get("reference")

        if object["object"][1] == "action":
            continue

        # calculate the score based on the 4 aspects used.
        score = 0
        score += (sentence_id - object["sentence_id"]) * distance_penalty
        if is_copula(object_to_resolve["object"][0]):
            if object_to_resolve["subject"] != object["subject"]:
                score += role_match_score
            if not object["subject"]:
                score += object_role_score
        else:
            if object_to_resolve["subject"] == object["subject"]:
                score += role_match_score
            if object["subject"]:
                score += subject_role_score

        results.append((object, score))

    return choose_best_fit(results)


def anaphora_resolution(
    sentences: List[Dict[str, Any]], manual_resolution: Dict[str, str] = {}
):
    """[Algorithm 11] Identify which objects need to be resolved and resolve
    them.

    Args:
        sentences (List[Dict[str,Any]]): All sentences and their extracted
            events.
        manual_resolution (Dict[str, str], optional): A mapping of manual
            resolution. Defaults to {}.

    Returns:
        List[Dict[str,Any]]: The input sentences but with references added.
    """
    for i, sentence in enumerate(sentences):
        for j, event in enumerate(sentence["events"]):
            for object in ["actor", "object"]:
                if not event[object]:
                    continue
                word = event[object]["object"][0]

                # Evaluate if the word needs to be resolved. by identifying if
                # it is a personal pronoun, a determiner or part of a
                # pre-defined list.
                if (
                    ("VNW" in word.tag_)
                    or ("LID" in word.tag_)
                    or (word.text in ["iemand", "iets"])
                ):
                    # Apply the manual resolution if possible.
                    if word.text in manual_resolution.keys():
                        sentences[i]["events"][j][object]["reference"] = (
                            manual_resolution[word.text]
                        )
                    else:
                        # if the word is a determiner use the action before
                        # this one as the reference.
                        if "LID" in word.tag_:
                            if j > 0:
                                sentences[i]["events"][j][object]["reference"] = (
                                    sentences[i]["events"][j - 1]["action"]
                                )
                            # If we have the first event of a sentence look for
                            # the last event of the previous sentence.
                            else:
                                if i > 0:
                                    index_last_event = (
                                        len(sentences[i - 1]["events"]) - 1
                                    )
                                    if index_last_event >= 0:
                                        sentences[i]["events"][j][object][
                                            "reference"
                                        ] = sentences[i - 1]["events"][
                                            index_last_event
                                        ][
                                            "action"
                                        ]
                        # In any other case we use the resolution algorithm
                        # above.
                        else:
                            sentences[i]["events"][j][object]["reference"] = (
                                find_reference(i, event[object], sentences)
                            )

    # If the reference we found has a reference itself we assign the reference
    # of the reference instead. This way a pronoun cannot refer to another
    # pronoun for example.
    for i, sentence in enumerate(sentences):
        for j, event in enumerate(sentence["events"]):
            for object in ["actor", "object"]:
                if not event[object]:
                    continue
                if sentences[i]["events"][j][object]["reference"]:
                    if sentences[i]["events"][j][object]["reference"].get("reference"):
                        sentences[i]["events"][j][object]["reference"] = sentences[i][
                            "events"
                        ][j][object]["reference"]["reference"]

    return sentences


def extract_events_for_article(document: str) -> Dict[str, Any]:
    """Given an Article of the law extract all events in that Article.

    Args:
        document (str): The content of the Article of the law.

    Returns:
        Dict[str, Any]: The extracted events form the Article.
    """
    # Create a parsing tree.
    parse_tree = nlp(document)

    # Split the document into sentences.
    sentences = [
        "".join(w.text_with_ws for w in sentence) for sentence in parse_tree.sents
    ]
    information_for_each_sentence = []

    # Extract element for each sentence and save all information
    for i, sentence in enumerate(sentences):
        sentence_tree = nlp(sentence)
        extracted_elements = sentence_decomposition(sentence_tree, i)
        sentence_information = {
            "sentence_id": i,
            "sentence": sentence,
            "sentence_tree": sentence_tree,
            "events": extracted_elements,
        }
        information_for_each_sentence.append(sentence_information)

    # Resolve references.
    final_sentences = anaphora_resolution(information_for_each_sentence)

    return final_sentences


def generate_chat_gpt_input(
    sentence_information: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """Print the event descriptions with the actor, action, and object text and
    references. This information is used as input to the improvement step done
    by ChatGPT.

    Args:
        sentence_information (List[Dict[str, Any]]): The processed sentences
        with all events.

    Returns:
        List[Dict[str,Any]]: The ChatGPT input.
    """
    final_output = []
    for sentence in sentence_information:
        # for each event extract the event description, the actor, action and
        # object text and the references
        for event in sentence["events"]:
            # get action text
            action = {"text": event["action"]["text"]}
            if event["actor"]:
                # get actor text and reference
                if event["actor"]["reference"]:
                    actor = {
                        "text": event["actor"]["text"],
                        "reference": event["actor"]["reference"]["text"],
                    }
                else:
                    actor = {"text": event["actor"]["text"], "reference": None}
            else:
                actor = None
            if event["object"]:
                # get object text and reference
                if event["object"]["reference"]:
                    object = {
                        "text": event["object"]["text"],
                        "reference": event["object"]["reference"]["text"],
                    }
                else:
                    object = {"text": event["object"]["text"], "reference": None}
            else:
                object = None

            object_reference = None
            actor_reference = None

            # Get the event description in the same we we do when we create the
            # flowchart, so excluding marks and replacing objects and actors
            # with their reference if there is one.

            if event["object"]:
                if event["object"]["reference"]:
                    object_reference = event["object"]["reference"]["text"]
                object_text = [
                    word
                    for word in event["object"]["relevant_words"]
                    if word.dep_ != "mark"
                ]
            else:
                object_text = []

            if event["actor"]:
                if event["actor"]["reference"]:
                    actor_reference = event["actor"]["reference"]["text"]
                actor_text = [
                    word
                    for word in event["actor"]["relevant_words"]
                    if word.dep_ != "mark"
                ]
            else:
                actor_text = []

            action_text = [
                word
                for word in event["action"]["relevant_words"]
                if word.dep_ != "mark"
            ]

            all_relevant_words = list(
                set(list(object_text) + list(actor_text) + list(action_text))
            )
            # make sure the words are all in the same order as they are in the original text
            all_relevant_words.sort(key=lambda word: word.i)

            # replace the object and actor with the word they refer to in the text.
            if object_reference:
                all_relevant_words[all_relevant_words.index(list(object_text)[0])] = (
                    object_reference
                )
            if actor_reference:
                all_relevant_words[all_relevant_words.index(list(actor_text)[0])] = (
                    actor_reference
                )

            # put everything in lowercase
            all_relevant_words = [
                word.text.lower() if not isinstance(word, str) else word
                for word in all_relevant_words
            ]
            event_text = " ".join(all_relevant_words)

            new_event = {
                "event_description": event_text,
                "action": action,
                "actor": actor,
                "object": object,
            }
            final_output.append(new_event)
    # print the output so it can be copied and fed to ChatGPT
    print(final_output)
    return final_output


def store_results_per_article(sentences: Dict, output_path: str) -> None:
    """Store intermediate output into a Word file for feedback.

    Args:
        sentences (Dict): The processed sentences with all relevant information
        output_path (str): the path where to save the output.
    """
    article_document = Document()

    for sentence in sentences:
        article_document.add_heading(f"Sentence {sentence['sentence_id']}", level=3)

        sentence_paragraph = article_document.add_paragraph()
        actor_words = []
        action_words = []
        object_words = []
        for event in sentence["events"]:
            if event["action"]:
                action_words += event["action"]["relevant_words"]
            if event["actor"]:
                actor_words += event["actor"]["relevant_words"]
            if event["object"]:
                object_words += event["object"]["relevant_words"]

        # Put action words in red, actor words in green and object words in
        # blue. All other words are simply printed in black.
        for word in sentence["sentence_tree"]:
            if word in action_words:
                sentence_paragraph.add_run(word.text_with_ws).font.color.rgb = RGBColor(
                    255, 0, 0
                )
            elif word in actor_words:
                sentence_paragraph.add_run(word.text_with_ws).font.color.rgb = RGBColor(
                    0, 255, 0
                )
            elif word in object_words:
                sentence_paragraph.add_run(word.text_with_ws).font.color.rgb = RGBColor(
                    0, 0, 255
                )
            else:
                sentence_paragraph.add_run(word.text_with_ws)

        # Below each sentence print an overview of all extracted events
        if not (
            (len(sentence["events"]) == 1) and (not sentence["events"][0]["action"])
        ):
            event_paragraph = article_document.add_paragraph()
            event_paragraph.add_run("Extracted events:")
        event_list = None
        for event in sentence["events"]:
            if event["action"]:
                event_list = article_document.add_paragraph(style="List Number")
                event_list.add_run("Action: " + event["action"]["text"])
            if event["actor"]:
                if not event_list:
                    event_list = article_document.add_paragraph(style="List Number")
                event_list.add_run(", actor: " + event["actor"]["text"])
                if event["actor"]["reference"]:
                    event_list.add_run(f" ({event['actor']['reference']['text']})")
            if event["object"]:
                if not event_list:
                    event_list = article_document.add_paragraph(style="List Number")
                event_list.add_run(", object: " + event["object"]["text"])
                if event["object"]["reference"]:
                    event_list.add_run(f" ({event['object']['reference']['text']})")

    article_document.save(output_path)


def find_action(relation: Token, event_list: List[Dict[str, any]]) -> Optional[int]:
    """Find the action belonging to the extracted relation.

    Args:
        relation (Token): the relation that should be related to an action
        event_list (List[Dict[str, any]]): a list of all events in the sentence
            containing the relation.

    Returns:
        Optional[int]: the index of the action belonging to the relation.
    """
    # Get a list of all actions.
    actions = [event["action"] for event in event_list]
    for k, action in enumerate(actions):
        if action:
            # We are looking for the actions that are the parent of the
            # relation token.
            if relation.head == action["object"][0]:
                return k


def marker_detection(
    sentence_information: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """[Algorithm 13] Detect one word markers in the text and attach them to
    the associated actions.

    Args:
        sentence_information (List[Dict[str, Any]]): The sentences to analyse

    Returns:
        List[Dict[str, Any]]: The input sentences with markers added to the
            actions.
    """
    for i, sentence in enumerate(sentence_information):
        # find all markers with the mark dependency
        marks = find_dependency("mark", sentence["sentence_tree"], all=True)
        for mark in marks:
            relevant_action = find_action(mark, sentence["events"])
            if relevant_action is not None:
                sentence_information[i]["events"][relevant_action]["action"][
                    "markers"
                ].append(mark.text.lower())
                # If the action is apart of a clausal complement we want to
                # indicate that the marker comes from a complementizer.
                if mark.head.dep_ == "ccomp":
                    sentence_information[i]["events"][relevant_action]["action"][
                        "marker_from_complementizer"
                    ].append(True)
                else:
                    sentence_information[i]["events"][relevant_action]["action"][
                        "marker_from_complementizer"
                    ].append(False)

        # Find all markers with the advmod dependency
        mods = find_dependency("advmod", sentence["sentence_tree"], all=True)
        for mod in mods:
            relevant_action = find_action(mod, sentence["events"])
            if relevant_action is not None:
                sentence_information[i]["events"][relevant_action]["action"][
                    "markers"
                ].append(mod.text.lower())
                sentence_information[i]["events"][relevant_action]["action"][
                    "marker_from_complementizer"
                ].append(False)

    return sentence_information


def detect_compound_indicators(
    sentence_information: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """[Algorithm 14] detect compound indicators of connections between actions
    using keyword lists.

    Args:
        sentence_information (List[Dict[str, Any]]): The sentences to analyse.

    Returns:
        List[Dict[str, Any]]: the input sentences markers added for the
            compound indicators.
    """
    # Define the stop word lists.
    conditional_indicators = [
        "als",
        "hetzij",
        "of",
        "indien",
        "geval",
        "anders",
        "optioneel",
        "wanneer",
    ]

    parallel_indicators = [
        "terwijl",
        "tussentijd",
        "parallel",
        "gelijktijdig",
        "ondertussen",
        "tegelijkertijd",
    ]

    sequence_indicators = [
        "dan",
        "na",
        "nadien",
        "daarna",
        "vervolgens",
        "basis",
        "gebaseerd",
        "dus",
    ]

    for i, sentence in enumerate(sentence_information):
        actions = [event["action"] for event in sentence["events"]]
        for j, action in enumerate(actions):
            if action:
                # Go through all words relevant to the action and check if
                # there are any matches with the words in the stop word lists.
                relevant_words = [
                    word.text.lower() for word in action["relevant_words"]
                ]

                # add c to the marker saved so we can distinguish what marker
                # came from where.
                if set(relevant_words).intersection(conditional_indicators):
                    sentence_information[i]["events"][j]["action"]["markers"].append(
                        "als c"
                    )
                    sentence_information[i]["events"][j]["action"][
                        "marker_from_complementizer"
                    ].append(True)

                elif set(relevant_words).intersection(sequence_indicators):
                    sentence_information[i]["events"][j]["action"]["markers"].append(
                        "dan c"
                    )
                    sentence_information[i]["events"][j]["action"][
                        "marker_from_complementizer"
                    ].append(False)

                elif set(relevant_words).intersection(parallel_indicators):
                    sentence_information[i]["events"][j]["action"]["markers"].append(
                        "terwijl c"
                    )
                    sentence_information[i]["events"][j]["action"][
                        "marker_from_complementizer"
                    ].append(False)

    return sentence_information


def determine_conjunct_events(sentence: Dict[str, Any], event_id: int) -> List[int]:
    """This function finds all conjunct events that are connected to the input
    event within the input sentence.

    Args:
        sentence (Dict[str, Any]): The input sentence where conjunctions can
            be found
        event_id (int): The id within the sentence of the event for which we
            want to find conjunct events

    Returns:
        List[int]: A list with the id's of the conjunct events.
    """
    conjunct_indeces = []
    sentence_tree = sentence["sentence_tree"]
    dependencies = find_dependency("conj", sentence_tree, all=True)
    for dependency in dependencies:
        for element in ["action", "actor", "object"]:
            if sentence["events"][event_id][element]:
                # if an element in the event of interest is connected to the
                # conjunction dependency we try to find the corresponding event.
                if (
                    sentence["events"][event_id][element]["object"][0]
                    == dependency.head
                ):
                    for k, event in enumerate(sentence["events"]):
                        if event[element]:
                            # In the corresponding event the same element should be
                            # the word associated with the conjunction dependency.
                            if event[element]["object"][0] == dependency:
                                conjunct_indeces.append(k)

    return conjunct_indeces


def add_implicit_markers(
    sentence_information: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """[Algorithm 15] Add implicit markers to the list of markers for each
    activity.

    Args:
        sentence_information (List[Dict[str, Any]]): All information relevant
            to each sentence in a document.

    Returns:
        List[Dict[str, Any]]: The input data with implicit markers added.
    """
    for i, sentence in enumerate(sentence_information):
        next_marker = None
        actions = [event["action"] for event in sentence["events"]]
        for j, action in enumerate(actions):
            # if we found an if-type statement before we add then then (dan) to
            # the next event as an if statement is logically followed by a then.
            if next_marker:
                sentence_information[i]["events"][j]["action"]["markers"].append(
                    next_marker
                )
                sentence_information[i]["events"][j]["action"][
                    "marker_from_complementizer"
                ].append(False)
                next_marker = None
            if sentence_information[i]["events"][j]["action"]:
                if sentence_information[i]["events"][j]["action"]["markers"] != []:
                    # see if there is an if-type statement using a stop word
                    # list of conditional indicators.
                    if set(
                        sentence_information[i]["events"][j]["action"]["markers"]
                    ).intersection(
                        {
                            "als",
                            "hetzij",
                            "of",
                            "indien",
                            "geval",
                            "terwijl",
                            "anders",
                            "optioneel",
                            "wanneer",
                            "als c",
                        }
                    ):
                        next_marker = "dan"
                    # If we find a sequential indicator (using a stop word list)
                    # then we make sure all actions of conjunct events get the
                    # same markers as this one.
                    if set(
                        sentence_information[i]["events"][j]["action"]["markers"]
                    ).intersection(
                        {
                            "dan",
                            "na",
                            "nadien",
                            "daarna",
                            "vervolgens",
                            "basis",
                            "gebaseerd",
                            "dus",
                            "dan c",
                        }
                    ):
                        conjunct_actions = determine_conjunct_events(sentence, j)
                        for conjunct_action in conjunct_actions:
                            if conjunct_action != j:
                                # find what markers this action has that are
                                # not already part of action of the conjunct
                                # event.
                                markers_to_add = [
                                    marker
                                    for marker in action["markers"]
                                    if marker
                                    not in sentence_information[i]["events"][
                                        conjunct_action
                                    ]["action"]["markers"]
                                ]
                                from_complementizer_to_add = [
                                    marker_from_complementizer
                                    for k, marker_from_complementizer in enumerate(
                                        action["marker_from_complementizer"]
                                    )
                                    if action["markers"][k]
                                    not in sentence_information[i]["events"][
                                        conjunct_action
                                    ]["action"]["markers"]
                                ]

                                sentence_information[i]["events"][conjunct_action][
                                    "action"
                                ]["markers"] += markers_to_add

                                sentence_information[i]["events"][conjunct_action][
                                    "action"
                                ][
                                    "marker_from_complementizer"
                                ] += from_complementizer_to_add

    return sentence_information


def correct_order(sentence_information: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """[Algorithm 16] Fix the ordering of the events in case there is an if
    clause following a main clause.

    Args:
        sentence_information (List[Dict[str, Any]]): All information relevant
            to each sentence in a document.

    Returns:
        List[Dict[str, Any]]: All information relevant to each sentence in a
            document with corrected event ordering.
    """
    conditional_indicators = [
        "als",
        "hetzij",
        "of",
        "indien",
        "geval",
        "anders",
        "optioneel",
        "wanneer",
        "als c",
    ]

    for i, sentence in enumerate(sentence_information):
        for j, event in enumerate(sentence["events"]):
            if event["action"]:
                if event["action"]["markers"] != []:
                    swap_events = False
                    # if we find a conditional indicator among the markers and
                    # this marker does not come from a complementizer we swap
                    # events
                    if set(event["action"]["markers"]).intersection(
                        conditional_indicators
                    ):
                        for indicator in conditional_indicators:
                            if indicator in event["action"]["markers"]:
                                if not event["action"]["marker_from_complementizer"][
                                    event["action"]["markers"].index(indicator)
                                ]:
                                    swap_events = True
                    if swap_events:
                        to_swap = sentence["events"][j - 1], sentence["events"][j]
                        (
                            sentence_information[i]["events"][j],
                            sentence_information[i]["events"][j - 1],
                        ) = to_swap

    return sentence_information


def find_action_with_object(
    sentence_information: List[Dict[str, Any]], object: Dict[str, Any]
) -> Dict[str, Any]:
    """Given an object  or actor of an action find the corresponding event.

    Args:
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in the document.
        object (Dict[str, Any]): The object or actor for which we have to find
            the event

    Returns:
        Dict[str, Any]: The event containing the object or actor.
    """
    for sentence in sentence_information:
        for event in sentence["events"]:
            if (event["object"] == object) or (event["actor"] == object):
                return event

    # If no event was found return None
    return None


def can_be_merged(
    event_1: Dict[str, Any], event_2: Dict[str, Any], ignore_negation_modifier: bool
) -> bool:
    """This function checks if two events can be merged based on 4 conditions
    (5 if ignore_negation_modifier is False): both have no or the same marker,
    one of the actions is weak, one of the actors is missing or requires a
    reference, one of the objects is missing or needs a reference(, and if
    ignore_negation_modifier is False we also check if both do or both do not
    have a negation modifier).

    Args:
        event_1 (Dict[str, Any]): One of the events that is checked for merging.
        event_2 (Dict[str, Any]): One of the events that is checked for merging.
        ignore_negation_modifier (bool): If True we do not look at the negation
            modifier.

    Returns:
        bool: True if the two actions can be merged according to the applied
            conditions and False otherwise.
    """
    relevant_words_event_1 = [
        word.text.lower() for word in event_1["action"]["relevant_words"]
    ]
    relevant_words_event_2 = [
        word.text.lower() for word in event_2["action"]["relevant_words"]
    ]

    # Check for negation modifiers if ignore_negation_modifier is False
    if not ignore_negation_modifier:
        event_1_negation = [
            word
            for word in relevant_words_event_1
            if word in ["niet", "geen", "nooit", "niemand", "nergens", "niets"]
        ]
        event_2_negation = [
            word
            for word in relevant_words_event_2
            if word in ["niet", "geen", "nooit", "niemand", "nergens", "niets"]
        ]
        if ((len(event_1_negation) < 0) and (len(event_2_negation) > 0)) or (
            (len(event_1_negation) > 0) and (len(event_2_negation) < 0)
        ):
            return False

    action_1 = event_1["action"]
    action_2 = event_2["action"]

    # Check if both have the same marker or no markers

    no_common_markers = (
        len([marker for marker in action_1["markers"] if marker in action_2["markers"]])
        == 0
    )
    at_least_one_marker = (len(action_1["markers"]) > 0) or (
        len(action_2["markers"]) > 0
    )

    if no_common_markers and at_least_one_marker:
        return False

    # check for weak verbs using a stop word list getting the lemma of the
    # verbs in the action.

    lemmas_verbs_event_1 = [
        word.lemma for word in event_1["action"]["relevant_words"] if "WW" in word.tag_
    ]
    lemmas_verbs_event_2 = [
        word.lemma for word in event_2["action"]["relevant_words"] if "WW" in word.tag_
    ]

    weak_verbs = [
        "zijn",
        "hebben",
        "doen",
        "berijken",
        "starten",
        "beginnen",
        "bestaan",
        "baseren",
    ]

    event_1_is_weak = set(lemmas_verbs_event_1).issubset(weak_verbs)
    event_2_is_weak = set(lemmas_verbs_event_2).issubset(weak_verbs)

    if (event_1_is_weak and event_2_is_weak) or (
        not event_1_is_weak and not event_2_is_weak
    ):
        return False

    # Check if actors are missing or need to be resolved

    event_1_no_actor = event_1["actor"] is None
    event_2_no_actor = event_2["actor"] is None

    if not event_1_no_actor:
        event_1_actor_needs_reference = (
            True if event_1.get("actor", {}).get("reference") else False
        )
    else:
        event_1_actor_needs_reference = False

    if not event_2_no_actor:
        event_2_actor_needs_reference = (
            True if event_2.get("actor", {}).get("reference") else False
        )
    else:
        event_2_actor_needs_reference = False

    # The line below  makes use of the fact that if event_1_no_actor is True
    # event_1_needs_reference is always False and the same for event 2.
    if (event_1_no_actor == event_1_actor_needs_reference) == (
        event_2_no_actor == event_2_actor_needs_reference
    ):
        return False

    # Check if objects are missing or need to be resolved

    event_1_no_object = event_1.get("object") is None
    event_2_no_object = event_2.get("object") is None

    if not event_1_no_object:
        event_1_object_needs_reference = (
            True if event_1.get("object", {}).get("reference") else False
        )
    else:
        event_1_object_needs_reference = False

    if not event_2_no_object:
        event_2_object_needs_reference = (
            True if event_2.get("object", {}).get("reference") else False
        )
    else:
        event_2_object_needs_reference = False

    if (event_1_no_object == event_1_object_needs_reference) == (
        event_2_no_object == event_2_object_needs_reference
    ):
        return False

    # If all above conditions hold and we did not find a rule that does not
    # apply we can return True

    return True


def merge_actions(
    event_1: Dict[str, Any], event_2: Dict[str, Any]
) -> Tuple[Dict[str, Any], bool]:
    """[Algorithm 18] Merge two events into one.

    Args:
        event_1 (Dict[str, Any]): The first event to merge.
        event_2 (Dict[str, Any]): The second event to merge.

    Returns:
        Tuple[Dict[str, Any], bool]: The new event and an indicator if the
            second event should be removed (True) or if the first event should
            be removed (False).
    """
    # Identify which event has the weak action assuming that one of the two
    # must have it based on the conditions stated in the can be merged function.
    lemmas_verbs_event_1 = [
        word.lemma for word in event_1["action"]["relevant_words"] if "WW" in word.tag_
    ]

    weak_verbs = [
        "zijn",
        "hebben",
        "doen",
        "berijken",
        "starten",
        "beginnen",
        "bestaan",
        "baseren",
    ]

    event_1_is_weak = set(lemmas_verbs_event_1).issubset(weak_verbs)

    if event_1_is_weak:
        weak_action = event_1
        strong_action = event_2
        remove_event_2 = False
    else:
        weak_action = event_2
        strong_action = event_1
        remove_event_2 = True

    # If the strong event has no actor or an actor that needs to be resolved
    # copy the one from the weak event if it exists.

    weak_no_actor = weak_action["actor"] is None
    strong_no_actor = strong_action["actor"] is None

    if not weak_no_actor:
        weak_actor_needs_reference = (
            True if weak_action.get("actor", {}).get("reference") else False
        )
    else:
        weak_actor_needs_reference = False

    if not strong_no_actor:
        strong_actor_needs_reference = (
            True if strong_action.get("actor", {}).get("reference") else False
        )
    else:
        strong_actor_needs_reference = False

    if (strong_no_actor and not weak_no_actor) or (
        strong_actor_needs_reference and not weak_actor_needs_reference
    ):
        strong_action["actor"] = weak_action["actor"]

    # If the strong event has no object or an object that needs to be resolved
    # copy the one from the weak event if it exists.

    weak_no_object = weak_action["object"] is None
    strong_no_object = strong_action["object"] is None

    if not weak_no_object:
        weak_object_needs_reference = (
            True if weak_action.get("object", {}).get("reference") else False
        )
    else:
        weak_object_needs_reference = False

    if not strong_no_object:
        strong_object_needs_reference = (
            True if strong_action.get("object", {}).get("reference") else False
        )
    else:
        strong_object_needs_reference = False

    if (strong_no_object and not weak_no_object) or (
        strong_object_needs_reference and not weak_object_needs_reference
    ):
        strong_action["object"] = weak_action["object"]

    return strong_action, remove_event_2


def get_event_from_action(
    sentence_information: List[Dict[str, Any]], action: Dict[str, Any]
) -> Dict[str, Any]:
    """Find an event given its action

    Args:
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in a document.
        action (Dict[str, Any]): The action for which we are trying to find the
            event.

    Returns:
        Dict[str, Any]: The event containing the input action.
    """
    for sentence in sentence_information:
        for event in sentence["events"]:
            if event["action"] == action:
                return event


def combine_actions(sentence_information: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """[Algorithm 17] Check if two actions can be combined when there is a
    reference relationship between them, and if so merge them or copy
    information from one to the other.

    Args:
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in the document.

    Returns:
        List[Dict[str, Any]]: The updated sentence information.
    """
    # Loop over list copies so when elements are removed the loop stays intact.
    for i, sentence in enumerate(sentence_information[:]):
        for j, event in enumerate(sentence["events"][:]):
            # Events can only be combined if one event is has a reference to an
            # element of another event. Hence we identify if there is another
            # event that we are referring to.
            referenced_action = None
            if event["actor"]:
                # check if the actor refers to an action
                if event["actor"]["reference"]:
                    if event["actor"]["reference"]["object"][1] == "action":
                        referenced_action = get_event_from_action(
                            sentence_information, event["actor"]["reference"]
                        )
            if not referenced_action:
                if event["object"]:
                    # Check if the object refers to an action, actor or object.
                    if event["object"]["reference"]:
                        if event["object"]["reference"]["object"][1] == "action":
                            referenced_action = get_event_from_action(
                                sentence_information, event["object"]["reference"]
                            )
                        else:
                            referenced_action = find_action_with_object(
                                sentence_information, event["object"]["reference"]
                            )

            if referenced_action:
                # store indexes to properly adapt the information if needed
                referenced_sentence_index = referenced_action["action"]["sentence_id"]
                referenced_action_index = sentence_information[
                    referenced_sentence_index
                ]["events"].index(referenced_action)
                if can_be_merged(event, referenced_action, True):
                    # merge the two events into one and remove the "weak" one.
                    merged_action, remove_reference = merge_actions(
                        event, referenced_action
                    )
                    if remove_reference:
                        sentence_information[i]["events"][j] = merged_action
                        sentence_information[referenced_sentence_index][
                            "events"
                        ].remove(referenced_action)
                    else:
                        sentence_information[referenced_sentence_index]["events"][
                            referenced_action_index
                        ] = merged_action
                        sentence_information[i]["events"].remove(event)
                elif can_be_merged(event, referenced_action, False):
                    # Copy actor and object information between events.
                    if event["actor"] and not referenced_action["actor"]:
                        referenced_action["actor"] = event["actor"]
                    elif not event["actor"] and referenced_action["actor"]:
                        event["actor"] = referenced_action["actor"]
                    if event["object"] and not referenced_action["object"]:
                        referenced_action["object"] = event["object"]
                    elif not event["object"] and referenced_action["object"]:
                        event["object"] = referenced_action["object"]

                    sentence_information[i]["events"][j] = event
                    sentence_information[referenced_sentence_index]["events"][
                        referenced_action_index
                    ] = referenced_action

    return sentence_information


def is_linkable(event_1: Dict[str, Any], event_2: Dict[str, Any]) -> bool:
    """This function determines if two actions are linkable based on 6 elements
    that need to match between the two: the copula specifier, the negation
    status, the actor, the object, the open clausal complement, and
    prepositional specifiers whose head word is "naar" or "over".

    Args:
        event_1 (Dict[str, Any]): The first event
        event_2 (Dict[str, Any]): The second event

    Returns:
        bool: True if the two events can be linked, False otherwise.
    """
    # check the copula specifier
    important_action_words_1 = event_1["action"]["relevant_words"]
    important_action_words_2 = event_2["action"]["relevant_words"]

    copula_specifier_1 = [
        word.head.text.lower()
        for word in important_action_words_1
        if word.dep_ == "cop"
    ]
    copula_specifier_2 = [
        word.head.text.lower()
        for word in important_action_words_2
        if word.dep_ == "cop"
    ]

    if copula_specifier_1 != copula_specifier_2:
        return False

    # Check the negation status (taken from the can_be_merged_function)
    event_1_negation = [
        word.text.lower()
        for word in important_action_words_1
        if word.text.lower() in ["niet", "geen", "nooit", "niemand", "nergens", "niets"]
    ]
    event_2_negation = [
        word.text.lower()
        for word in important_action_words_2
        if word.text.lower() in ["niet", "geen", "nooit", "niemand", "nergens", "niets"]
    ]
    if ((len(event_1_negation) < 0) and (len(event_2_negation) > 0)) or (
        (len(event_1_negation) > 0) and (len(event_2_negation) < 0)
    ):
        return False

    # Check if they have the same actor
    if event_1["actor"]:
        actor_1 = event_1["actor"]["text"]
    else:
        actor_1 = None
    if event_2["actor"]:
        actor_2 = event_2["actor"]["text"]
    else:
        actor_2 = None

    if actor_1 != actor_2:
        return False

    # Check if they have the same object
    if event_1["object"]:
        object_1 = event_1["object"]["text"]
    else:
        object_1 = None
    if event_2["object"]:
        object_2 = event_2["object"]["text"]
    else:
        object_2 = None

    if object_1 != object_2:
        return False

    # Find out if the xcomp elements match
    xcomp_event_1 = [word for word in important_action_words_1 if word.dep_ == "xcomp"]
    xcomp_event_2 = [word for word in important_action_words_2 if word.dep_ == "xcomp"]

    if xcomp_event_1 != xcomp_event_2:
        return False

    # Find prepositional specifiers whose head word is to or about using subtrees
    prep_specifier_subtree_1 = [
        word.head.subtree
        for word in important_action_words_1
        if word.text.lower() in ["naar", "over"]
    ]
    prep_specifier_subtree_2 = [
        word.head.subtree
        for word in important_action_words_2
        if word.text.lower() in ["naar", "over"]
    ]

    if len(prep_specifier_subtree_1) > 0:
        prep_specifier_1 = [word.text.lower() for word in prep_specifier_subtree_1[0]]
    else:
        prep_specifier_1 = None
    if len(prep_specifier_subtree_2) > 0:
        prep_specifier_2 = [word.text.lower() for word in prep_specifier_subtree_2[0]]
    else:
        prep_specifier_2 = None

    if prep_specifier_1 != prep_specifier_2:
        return False

    # If no contradictions to the conditions are found we return True
    return True


def event_contains_forward_links(event: Dict[str, Any]) -> bool:
    """Check if an event contains forward links based on a stop word list.

    Args:
        event (Dict[str, Any]): The event for which we have to check the possible
            link.

    Returns:
        bool: True if the event action contains a forward link and false otherwise.
    """
    relevant_words = event["action"]["relevant_words"]

    forward_link_words = ["slot", "slotte", "eindelijk", "uiteindelijk"]

    for word in relevant_words:
        if word in forward_link_words:
            return True

    return False


def event_contains_loop_links(event: Dict[str, Any]) -> bool:
    """Check if an event contains loop links based on a stop word list.

    Args:
        event (Dict[str, Any]): The event for which we have to check the possible
            link.

    Returns:
        bool: True if the event action contains a loop link and False otherwise.
    """
    relevant_words = event["action"]["relevant_words"]

    loop_link_words = [
        "volgende",
        "vervolgens",
        "daarna",
        "nadien",
        "hierna",
        "verder",
        "weer",
        "opnieuw",
        "nogmaals",
        "wederom",
        "terug",
        "terugkeren",
    ]

    for word in relevant_words:
        if word in loop_link_words:
            return True

    return False


def determine_link_type(
    source_event: Dict[str, Any],
    target_event: Dict[str, Any],
    sentence_information: List[Dict[str, Any]],
) -> str:
    """[Algorithm 20] This function determines the type of link between two
    events.

    Args:
        source_event (Dict[str, Any]): The event that is the source of the link
        target_event (Dict[str, Any]): The event that is the target of the link
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in the document.
    Returns:
        str: The type of link ("forward", "jump", "loop" or None)
    """
    conditional_indicators = [
        "als",
        "hetzij",
        "of",
        "indien",
        "geval",
        "anders",
        "optioneel",
        "wanneer",
        "als c",
    ]
    if len(source_event["action"]["markers"]) > 0:
        # if we find conditional markers we check for forward and jump links.
        if set(source_event["action"]["markers"]).intersection(conditional_indicators):
            if event_contains_forward_links(source_event):
                return "forward"

            target_sentence_id = target_event["action"]["sentence_id"]
            target_sentence = sentence_information[target_sentence_id]
            target_tree = target_sentence["sentence_tree"]

            dependencies = find_dependency("conj", target_tree, all=True)

            # If we find a conjunction relationship where the conjunction is
            # caused by an or-connection (the coordinating conjunction (cc) if
            # 'of') we have a jump loop
            for dependency in dependencies:
                for element in ["actor", "object"]:
                    if target_event[element]:
                        target_element = target_event.get(element, {}).get(
                            "object", [None]
                        )[0]
                        # Check ik the conjunction is with an element of this
                        # event.
                        if target_element == dependency.head:
                            for event in target_sentence["events"]:
                                if event[element]:
                                    # find the event with the conjunction we found
                                    conjunct_element = event.get(element, {}).get(
                                        "object", [None]
                                    )[0]
                                    if conjunct_element == dependency:
                                        # check if the conjunction is caused by
                                        # an if-relationship.
                                        ccs = find_dependency(
                                            "cc", target_tree, all=True
                                        )
                                        for cc in ccs:
                                            if (
                                                cc.head
                                                in [target_element, conjunct_element]
                                            ) and (cc.text.lower() == "of"):
                                                return "jump"
        else:
            if event_contains_loop_links(source_event):
                return "loop"
    else:
        if event_contains_loop_links(source_event):
            return "loop"

    return None


def determine_inter_action_links(
    sentence_information: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """[Algorithm 19] Identify links between actions

    Args:
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in the document.

    Returns:
        List[Dict[str, Any]]: The updated sentence information.
    """
    events = [
        event
        for sentence in sentence_information
        for event in sentence["events"]
        if event["action"]
    ]

    # Go over all event combinations and link the two events together if they
    # are linkable.
    for event_1, event_2 in combinations(events, 2):
        if is_linkable(event_1, event_2):
            event_1_sentence_id = event_1["action"]["sentence_id"]
            event_1_index = sentence_information[event_1_sentence_id]["events"].index(
                event_1
            )
            event_1["link"] = event_2
            link_type = determine_link_type(event_1, event_2, sentence_information)
            event_1["link_type"] = link_type
            sentence_information[event_1_sentence_id]["events"][event_1_index] = event_1

    return sentence_information


def determine_conjoined_elements(
    sentence: Dict[str, Any], event: Dict[str, Any], ignore_events: List[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str]]:
    """This function finds all events that have a conjunct element with the
    main event.

    Args:
        sentence (Dict[str, Any]): The information for the sentence containing
            the event.
        event (Dict[str, Any]): The event for which to find conjunctions.

    Returns:
        List[Tuple[Dict[str, Any], str]]: A list of conjunct events.
    """
    conjoined_result = []
    dependencies = find_dependency("conj", sentence["sentence_tree"], all=True)

    ignore_events.append(event)

    for dependency in dependencies:
        # check for conjunctions for each element
        for element in ["action", "actor", "object"]:
            if event[element]:
                target_element = event[element]["object"][0]
                if (target_element == dependency.head) or (
                    target_element == dependency
                ):
                    # Find the conjunct events by looking at all events that
                    # have the element we need.
                    for conjunct_event in [
                        sentence_event
                        for sentence_event in sentence["events"]
                        if sentence_event[element]
                    ]:
                        if conjunct_event not in ignore_events:
                            conjunct_element = conjunct_event.get(element, {}).get(
                                "object", [None]
                            )[0]
                            # If the event is conjunct to the main event we find
                            # the coordinating conjunction and store a combination
                            # of the two.
                            if (conjunct_element == dependency) or (
                                conjunct_element == dependency.head
                            ):
                                ccs = find_dependency(
                                    "cc", sentence["sentence_tree"], all=True
                                )
                                for cc in ccs:
                                    # Some words are seen as coordinating
                                    # conjunctions while they should not be
                                    # these are stored here to be skipped.
                                    if cc.text.lower() in ["a", "doch", "maar", "e."]:
                                        continue
                                    if cc.head in [target_element, conjunct_element]:
                                        # We keep track of all event that have
                                        # already been added to prevent infinite
                                        # loops.
                                        ignore_events.append(conjunct_event)
                                        conjoined_result.append(
                                            (conjunct_event, element, cc)
                                        )
                                        conjoined_result += (
                                            determine_conjoined_elements(
                                                sentence, conjunct_event, ignore_events
                                            )
                                        )

    return conjoined_result


def build_gateway(
    last_actions: List[Dict[str, Any]],
    open_split: List[Dict[str, Any]],
    conjoined_result: List[Tuple[Dict[str, Any], str]],
    event: Dict[str, Any],
    flow: List[Dict[str, Any]],
    type: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    """This function builds a split in the flow where two events will follow
    from the previous one.

    Args:
        last_actions (List[Dict[str, Any]]): All previous actions
        open_split (List[Dict[str, Any]]): The last actions in an open split
        conjoined_result (List[Tuple[Dict[str, Any], str]]): The event
            following the split
        event (Dict[str, Any]): Another event following the split
        flow (List[Dict[str, Any]]): All flow elements found so far
        type (str): The type of split being created ("of" or "en")

    Returns:
        Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
            The updated, flow, last_actions, and open_split lists.
    """
    # If there is a previous event this is our starting point
    if len(last_actions) > 0:
        from_event = last_actions[-1]
    else:
        from_event = None

    flow_element = {
        "from_event": from_event,
        "to_event": event,
        "type": type,
        "split": True,
    }
    flow.append(flow_element)
    last_actions.append(event)
    open_split.append(event)

    # make a flow from the starting point to each of the events that are a part
    # of this split and add the type of the split tot the flows.
    for conjoined in conjoined_result:
        flow_element = {
            "from_event": from_event,
            "to_event": conjoined[0],
            "type": type,
            "split": True,
        }
        flow.append(flow_element)
        last_actions.append(conjoined[0])
        # Add the events to the open_split
        open_split.append(conjoined[0])

    return flow, last_actions, open_split


def build_join(
    flow: List[Dict[str, Any]],
    last_actions: List[Dict[str, Any]],
    open_split: List[Dict[str, Any]],
    dummy_id: int,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
    """This functions joins the current split into a dummy node.

    Args:
        flow (List[Dict[str, Any]]): A list of all flow elements found so far.
        last_actions (List[Dict[str, Any]]): A list of previous events.
        open_split (List[Dict[str, Any]]): A list of event in the open split.
        dummy_id (int): ID for the dummy to distinguish them later.

    Returns:
        Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
            The updated flow, last_actions and open_split lists.
    """
    # For each event in the open split create a flow to the same dummy join node.
    for split_event in open_split:
        new_flow_element = {
            "from_event": split_event,
            "to_event": {"object": f"Dummy Node join {dummy_id}"},
            "type": "join",
            "split": False,
        }
        flow.append(new_flow_element)

    last_actions.append({"object": f"Dummy Node join {dummy_id}"})

    open_split = []

    return flow, last_actions, open_split


def handle_single_action(
    flow: List[Dict[str, Any]],
    event: Dict[str, Any],
    last_actions: List[Dict[str, Any]],
    open_split: List[Dict[str, Any]],
    dummy_id: int,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]], int]:
    """[Algorithm 22] Add a single action to the flow.

    Args:
        flow (List[Dict[str, Any]]): A list of all flow elements found so far.
        event (Dict[str, Any]): The event to add.
        last_actions (List[Dict[str, Any]]): A list of previous events.
        open_split (List[Dict[str, Any]]): A list of event in the open split.
        dummy_id (int): ID for the dummy to distinguish them later.

    Returns:
        Tuple[List[Dict[str, Any]], List[Dict[str, Any]], List[Dict[str, Any]]]:
            The updated flow, last_actions, open_split lists, and dummy_id.
    """
    # Define the stop word lists.
    conditional_indicators = [
        "als",
        "hetzij",
        "of",
        "indien",
        "geval",
        "anders",
        "optioneel",
        "wanneer",
        "als c",
    ]

    parallel_indicators = [
        "terwijl",
        "tussentijd",
        "parallel",
        "gelijktijdig",
        "ondertussen",
        "tegelijkertijd",
        "terwijl c",
    ]

    sequence_indicators = [
        "dan",
        "na",
        "nadien",
        "daarna",
        "vervolgens",
        "basis",
        "gebaseerd",
        "dus",
        "dan c",
    ]
    # if we find sequence_indicators or conditional_indicators we close the
    # current split.
    if event["action"]["markers"] != []:
        if (set(event["action"]["markers"]).intersection(sequence_indicators)) or (
            set(event["action"]["markers"]).intersection(conditional_indicators)
        ):
            if open_split:
                flow, last_actions, open_split = build_join(
                    flow, last_actions, open_split, dummy_id
                )
                dummy_id += 1

    if event["action"]["markers"] != []:
        # if we find parallel indicators we create an and-split by putting this
        # event parallel to the last event added to the flowchart.
        if set(event["action"]["markers"]).intersection(parallel_indicators):
            if last_actions:
                flow_from = flow[-1]["from_event"]
                flow_to = flow[-1]["to_event"]
            else:
                flow_element = {
                    "from_event": None,
                    "to_event": {"object": f"Dummy Node split {dummy_id}"},
                    "type": "en",
                    "split": False,
                }
                flow.append(flow_element)
                flow_from = {"object": f"Dummy Node split {dummy_id}"}
                flow_to = None
                dummy_id += 1

            new_flow = {
                "from_event": flow_from,
                "to_event": event,
                "type": "en",
                "split": True,
            }
            flow[-1]["type"] = "en"
            flow[-1]["split"] = True
            open_split = [flow_to, event] if flow_to else [event]
            flow.append(new_flow)
            last_actions.append(event)

        else:
            # otherwise we add the event sequentially. Either behind a start
            # node if it is the first event or behind the previous node.
            if len(last_actions) > 0:
                new_flow = {
                    "from_event": last_actions[-1],
                    "to_event": event,
                    "type": None,
                    "split": False,
                }
                flow.append(new_flow)
                if last_actions[-1] in open_split:
                    replace_index = open_split.index(last_actions[-1])
                    open_split[replace_index] = event
                last_actions.append(event)
            else:
                new_flow = {
                    "from_event": "start",
                    "to_event": event,
                    "type": None,
                    "split": False,
                }
                flow.append(new_flow)
                last_actions.append(event)
    else:
        if len(last_actions) > 0:
            new_flow = {
                "from_event": last_actions[-1],
                "to_event": event,
                "type": None,
                "split": False,
            }
            flow.append(new_flow)
            if last_actions[-1] in open_split:
                replace_index = open_split.index(last_actions[-1])
                open_split[replace_index] = event
            last_actions.append(event)
        else:
            new_flow = {
                "from_event": "start",
                "to_event": event,
                "type": None,
                "split": False,
            }
            flow.append(new_flow)
            last_actions.append(event)

    return flow, last_actions, open_split, dummy_id


def build_flows(sentence_information: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """[Algorithm 21] This function identifies the flow between all events in a
    document.

    Args:
        sentence_information (List[Dict[str, Any]]): A list with information on
            all sentences in the document.

    Returns:
        List[Dict[str, Any]]: A list of all Flows that will end up in the
            flowchart.
    """
    # 'dan' is included as the construction 'dan wel' is an or indicator
    or_indicators = ["of", "hetzij", "noch", "respectievelijk", "dan"]
    and_indicators = ["en", "alsmede", "onderscheidenlijk", "zowel", "tot"]

    last_actions = []
    open_split = []
    flow = []
    processed_events = []
    dummy_id = 0
    for i, sentence in enumerate(sentence_information):
        for j, event in enumerate(sentence["events"]):
            if event not in processed_events:
                # prevent events being processed twice.
                processed_events.append(event)
                if event["action"]:
                    # if we encounter a jump link we start a new part of the flow.
                    if event["link_type"] == "jump":
                        last_actions = []
                        open_split = []
                        new_flow = {
                            "from_event": None,
                            "to_event": event["link"],
                            "type": None,
                            "split": False,
                        }
                        flow.append(new_flow)
                        last_actions.append(event["link"])

                    # Find all conjunct events and the corresponding coordinating
                    # conjunctions.
                    conjoined_result = determine_conjoined_elements(sentence, event, [])
                    if len(conjoined_result) == 0:
                        # if no conjunctions are found, handle single action
                        flow, last_actions, open_split, dummy_id = handle_single_action(
                            flow, event, last_actions, open_split, dummy_id
                        )
                    else:
                        coordinating_conjunctions = [
                            conjoined_element[2].text.lower()
                            for conjoined_element in conjoined_result
                        ]

                        # If there are any coordinating conjunctions that are
                        # not either in the or_indicators or and_indicators we
                        # print this coordinating conjunction in an error
                        # message so it can be added to one of the lists or
                        # excluded as coordinating conjunction.
                        if any(
                            (coordinating_conj not in or_indicators)
                            and (coordinating_conj not in and_indicators)
                            for coordinating_conj in coordinating_conjunctions
                        ):
                            missing_words = [
                                coordinating_conj
                                for coordinating_conj in coordinating_conjunctions
                                if (coordinating_conj not in or_indicators)
                                and (coordinating_conj not in and_indicators)
                            ]
                            raise ValueError(
                                f"Connection word(s) {missing_words} has not been added to a list yet, please add it and re-run the code. Sentence: {sentence['sentence']}"
                            )

                        # If we only have or_indicators we make a or-split
                        # ('of' in Dutch) of all conjoined events.
                        if all(
                            coordinating_conj in or_indicators
                            for coordinating_conj in coordinating_conjunctions
                        ):
                            if last_actions != []:
                                dummy_flow = {
                                    "from_event": last_actions[-1],
                                    "to_event": {
                                        "object": f"Dummy Node split {dummy_id}"
                                    },
                                    "type": "of",
                                    "split": False,
                                }
                            else:
                                dummy_flow = {
                                    "from_event": "start",
                                    "to_event": {
                                        "object": f"Dummy Node split {dummy_id}"
                                    },
                                    "type": "of",
                                    "split": False,
                                }
                            flow.append(dummy_flow)
                            last_actions.append(
                                {"object": f"Dummy Node split {dummy_id}"}
                            )
                            dummy_id += 1
                            flow, last_actions, open_split = build_gateway(
                                last_actions,
                                open_split,
                                conjoined_result,
                                event,
                                flow,
                                "of",
                            )
                            # we immediately join the split
                            flow, last_actions, open_split = build_join(
                                flow,
                                last_actions,
                                open_split,
                                dummy_id,
                            )
                            dummy_id += 1
                            processed_events += [
                                conjoined[0] for conjoined in conjoined_result
                            ]

                        elif all(
                            coordinating_conj in and_indicators
                            for coordinating_conj in coordinating_conjunctions
                        ):
                            # If we only have and_indicators and the split is
                            # based on the actors we make a and-split ('en' in
                            # Dutch) of all conjoined events.
                            if all(
                                conjoined[1] == "actor"
                                for conjoined in conjoined_result
                            ):
                                if last_actions != []:
                                    dummy_flow = {
                                        "from_event": last_actions[-1],
                                        "to_event": {
                                            "object": f"Dummy Node split {dummy_id}"
                                        },
                                        "type": "en",
                                        "split": False,
                                    }
                                else:
                                    dummy_flow = {
                                        "from_event": "start",
                                        "to_event": {
                                            "object": f"Dummy Node split {dummy_id}"
                                        },
                                        "type": "en",
                                        "split": False,
                                    }
                                flow.append(dummy_flow)
                                last_actions.append(
                                    {"object": f"Dummy Node split {dummy_id}"}
                                )
                                dummy_id += 1
                                flow, last_actions, open_split = build_gateway(
                                    last_actions,
                                    open_split,
                                    conjoined_result,
                                    event,
                                    flow,
                                    "en",
                                )
                                # we immediately join the split
                                flow, last_actions, open_split = build_join(
                                    flow,
                                    last_actions,
                                    open_split,
                                    dummy_id,
                                )
                                dummy_id += 1
                                processed_events += [
                                    conjoined[0] for conjoined in conjoined_result
                                ]
                            # if we only have and_indicators but the conjunction
                            # is not based on the actor we simply close any open
                            # split and handle single action.
                            else:
                                if open_split:
                                    flow, last_actions, open_split = build_join(
                                        flow,
                                        last_actions,
                                        open_split,
                                        dummy_id,
                                    )
                                    dummy_id += 1
                                flow, last_actions, open_split, dummy_id = (
                                    handle_single_action(
                                        flow, event, last_actions, open_split, dummy_id
                                    )
                                )
                        # If there are both or_indicators and and_indicators
                        # we make an en/of split. Indicating any number of
                        # paths can be taken.
                        else:
                            if last_actions != []:
                                dummy_flow = {
                                    "from_event": last_actions[-1],
                                    "to_event": {
                                        "object": f"Dummy Node split {dummy_id}"
                                    },
                                    "type": "en/of",
                                    "split": False,
                                }
                            else:
                                dummy_flow = {
                                    "from_event": "start",
                                    "to_event": {
                                        "object": f"Dummy Node split {dummy_id}"
                                    },
                                    "type": "en/of",
                                    "split": False,
                                }
                            flow.append(dummy_flow)
                            last_actions.append(
                                {"object": f"Dummy Node split {dummy_id}"}
                            )
                            dummy_id += 1
                            flow, last_actions, open_split = build_gateway(
                                last_actions,
                                open_split,
                                conjoined_result,
                                event,
                                flow,
                                "en/of",
                            )
                            # We immediately join the split
                            flow, last_actions, open_split = build_join(
                                flow,
                                last_actions,
                                open_split,
                                dummy_id,
                            )
                            dummy_id += 1
                            processed_events += [
                                conjoined[0] for conjoined in conjoined_result
                            ]
    return flow


def generate_chat_gpt_output(
    flows: List[Dict[str, Any]],
    new_event_list: List[Dict[str, Any]],
    original_event_list: List[Dict[str, Any]],
) -> None:
    """Create the flowchart representation that ChatGPT needs to improve in
    its final task using the corrected event descriptions created by ChatGPT.

    Args:
        flows (List[Dict[str, Any]]): All flows extracted by the algorithm.
        new_event_list (List[Dict[str, Any]]): The corrected event list created
            by ChatGPT.
        original_event_list (List[Dict[str, Any]]): The original event list the
            flows are based on.
    """
    final_output = []
    for flow in flows:
        from_event_text = None
        to_event_text = None

        # find text of the from event node
        if flow["from_event"] == "start":
            from_event_text = "start"
        elif flow["from_event"] is None:
            from_event_text = " "
        elif flow["from_event"]["object"]:
            if "Dummy Node split" in flow["from_event"]["object"]:
                # include dummy id as ChatGPT should also include them
                dummy_id = re.findall(r"\d+", flow["from_event"]["object"])
                from_event_text = flow["type"] + f" {dummy_id[0]}"
            elif "Dummy Node join" in flow["from_event"]["object"]:
                dummy_id = re.findall(r"\d+", flow["from_event"]["object"])
                from_event_text = f"join {dummy_id[0]}"
        if not from_event_text:
            # If the event is not start None or a Dummy node it should be in
            # the original event list
            from_event_id = original_event_list.index(flow["from_event"])
            # We find the corresponding corrected event description created by
            # ChatGPT
            from_event_text = [
                event
                for event in new_event_list
                if event["original_event_id"] == (from_event_id + 1)
            ][0]["event_description"]

        # Similarly find the to event text
        if flow["to_event"]["object"]:
            if "Dummy Node split" in flow["to_event"]["object"]:
                dummy_id = re.findall(r"\d+", flow["to_event"]["object"])
                to_event_text = flow["type"] + f" {dummy_id[0]}"
            elif "Dummy Node join" in flow["to_event"]["object"]:
                dummy_id = re.findall(r"\d+", flow["to_event"]["object"])
                to_event_text = f"join {dummy_id[0]}"
        if not to_event_text:
            to_event_id = original_event_list.index(flow["to_event"])
            to_event_text = [
                event
                for event in new_event_list
                if event["original_event_id"] == (to_event_id + 1)
            ][0]["event_description"]

        # Define the flow based on the text only.
        new_flow = {"from_event": from_event_text, "to_event": to_event_text}
        final_output.append(new_flow)

    print(final_output)


def create_node_from_event(
    event: Dict[str, Any], flow: Dict[str, Any], node_id: int, graph: graphviz.Digraph
) -> graphviz.Digraph:
    """Create a node to put into the flowchart by extracting the relevant
    information from the event and the flow.

    Args:
        event (Dict[str, Any]): The event for which we create a node.
        flow (Dict[str, Any]): The flow showing the split type if we create a
            dummy node.
        node_id (int): The id of the node in the graph.
        graph (graphviz.Digraph): The flowchart.

    Returns:
        graphviz.Digraph: The flowchart with a node added.
    """
    object_reference = None
    actor_reference = None

    # get the event description
    if event["object"]:
        if "Dummy Node split" in event["object"]:
            graph.node(str(node_id), flow["type"])
            return graph
        elif "Dummy Node join" in event["object"]:
            graph.node(str(node_id), "Join")
            return graph

    # get all relevant words for the object
    if event["object"]:
        if event["object"]["reference"]:
            object_reference = event["object"]["reference"]["text"]
        object_text = [
            word for word in event["object"]["relevant_words"] if word.dep_ != "mark"
        ]
    else:
        object_text = []

    # get all relevant words for the actor
    if event["actor"]:
        if event["actor"]["reference"]:
            actor_reference = event["actor"]["reference"]["text"]
        actor_text = [
            word for word in event["actor"]["relevant_words"] if word.dep_ != "mark"
        ]
    else:
        actor_text = []

    # get all relevant words for the action
    action_text = [
        word for word in event["action"]["relevant_words"] if word.dep_ != "mark"
    ]

    all_relevant_words = list(
        set(list(object_text) + list(actor_text) + list(action_text))
    )
    # make sure all words are in the same order as they are in the original text
    all_relevant_words.sort(key=lambda word: word.i)

    # replace the actor and object with their reference if they have one.
    if object_reference:
        all_relevant_words[all_relevant_words.index(list(object_text)[0])] = (
            object_reference
        )
    if actor_reference:
        all_relevant_words[all_relevant_words.index(list(actor_text)[0])] = (
            actor_reference
        )

    all_relevant_words = [
        word.text.lower() if not isinstance(word, str) else word
        for word in all_relevant_words
    ]
    event_text = " ".join(all_relevant_words)

    graph.node(str(node_id), event_text)

    return graph


def create_flow_diagram(
    flow: List[Dict[str, Any]], article: str, output_path: str
) -> None:
    """Create the flow diagram image from the flows.

    Args:
        flow (List[Dict[str, Any]]): THe flows that should be made into a
            flowchart.
        article (str): The name of the article the flowchart is for.
        output_path (str): The path the flowchart should be saved to.
    """
    graph = graphviz.Digraph("article_flow", comment="Article of Interest")
    graph.node("start", "Start")
    event_node_labels = []
    processed_events = []
    for i, flow_element in enumerate(flow):
        # If the node for the event already exists find the node, if the node
        # for the event does not exist make the node.
        to_event = flow_element["to_event"]
        if to_event in processed_events:
            to_node_label = event_node_labels[processed_events.index(to_event)]
        else:
            graph = create_node_from_event(to_event, flow_element, i, graph)
            to_node_label = str(i)
            event_node_labels.append(str(i))
            processed_events.append(to_event)

        # Find the from node and create an edge between the from and to node.
        if flow_element["from_event"]:
            if flow_element["from_event"] == "start":
                graph.edge("start", to_node_label)
            else:
                from_event = flow_element["from_event"]
                from_node_label = event_node_labels[processed_events.index(from_event)]
                graph.edge(from_node_label, to_node_label)

    graph.render(
        os.path.join(output_path, article.replace(":", ".")), format="png", view=False
    )


if __name__ == "__main__":
    path_to_data = os.path.join(os.path.abspath(__file__), "../../data/")
    document_data = load_data(
        os.path.join(
            path_to_data,
            "BWBR0001903_BoekEerste-geldend_van_01-10-2023_tm_heden_zichtdatum_30-10-2023.txt",
        )
    )

    document_data["sentence_information"] = document_data["content"].apply(
        extract_events_for_article
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        marker_detection
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        detect_compound_indicators
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        add_implicit_markers
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        correct_order
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        combine_actions
    )

    document_data["sentence_information"] = document_data["sentence_information"].apply(
        determine_inter_action_links
    )

    document_data["flows"] = document_data["sentence_information"].apply(build_flows)

    path_to_output = os.path.join(os.path.abspath(__file__), "../../flows/")

    document_data.apply(
        lambda row: create_flow_diagram(row["flows"], row["article"], path_to_output),
        axis=1,
    )

    for index, row in document_data.iterrows():
        store_results_per_article(
            row["sentence_information"],
            f'extracted_events/Boek_1_{row["article"].replace(":", ".")}_events.docx',
        )
